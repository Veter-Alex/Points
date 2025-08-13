
"""
Модуль для создания Word-отчета по точкам наблюдения.

Содержит функцию create_word_report, которая формирует отчет с группировкой точек по странам и регионам.
Все функции снабжены подробными комментариями и докстрингами согласно лучшим практикам.
"""

import os
from typing import List, Dict

from models.points import PointRecord


def create_word_report(
    points_folder: List[PointRecord],
    wrong_city_data_folder: List[str],
    report_path: str,
    log_message=None
) -> bool:
    """
    Создает Word-отчет с точками, сгруппированными по странам и областям.

    Сначала выводятся точки на территории России (разные области), затем Украины, затем остальных стран.
    В конце добавляется информация о не найденных городах и оригинальные тексты точек.

    Args:
        points_folder (List[PointRecord]): Список точек для включения в отчет.
        wrong_city_data_folder (List[str]): Список городов, не найденных в city.txt.
        report_path (str): Путь для сохранения Word-файла.
        log_message (callable, optional): Функция для логирования.

    Returns:
        bool: True если отчет успешно создан, иначе False.
    """

    try:
        # Импортируем необходимые классы для работы с Word
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        from docx.shared import Pt

        # Создаем новый документ Word и настраиваем стиль
        doc = Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(14)
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # --- Группировка точек по странам ---
        russia_points = []
        ukraine_points = []
        other_points = []
        all_original_texts = []
        seen_coordinates = set()
        for point in points_folder:
            # Уникальность по координатам (округление до 6 знаков)
            coord_key = (round(point.latitude, 6), round(point.longitude, 6))
            if coord_key not in seen_coordinates:
                seen_coordinates.add(coord_key)
                if point.original_text:
                    all_original_texts.append(point.original_text)
            # Группируем по стране
            if point.country == "Russia":
                russia_points.append(point)
            elif point.country == "Ukraine":
                ukraine_points.append(point)
            else:
                other_points.append(point)

        def extract_region(area_desc: str) -> str:
            """
            Извлекает название региона из описания района.
            Если не найдено — возвращает 'Неизвестная область'.
            """
            if not area_desc:
                return "Неизвестная область"
            import re
            pattern = r"на территории (.+?)(?:,|$)"
            match = re.search(pattern, area_desc)
            if match:
                return match.group(1).strip()
            return "Неизвестная область"

        def add_region_header(text: str):
            """
            Добавляет заголовок региона в документ с курсивом.
            """
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            run.italic = True
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return p

        def add_normal_paragraph(text: str):
            """
            Добавляет обычный абзац с текстом в документ.
            """
            p = doc.add_paragraph(text)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
            return p

        def format_original_text(text: str) -> str:
            """
            Форматирует оригинальный текст точки: удаляет пустые строки и пробелы.
            """
            if not text:
                return ""
            lines = text.split("\n")
            formatted_lines = []
            for line in lines:
                clean_line = line.strip()
                if clean_line:
                    formatted_lines.append(clean_line)
            return "\n".join(formatted_lines)

        # --- Группировка точек России по регионам ---
        russia_by_region: Dict[str, List[PointRecord]] = {}
        for point in russia_points:
            region = (
                point.region_desc
                if point.region_desc
                else extract_region(point.area_desc or "")
            )
            if region not in russia_by_region:
                russia_by_region[region] = []
            russia_by_region[region].append(point)

        def get_unique_points(points_list):
            """
            Возвращает уникальные точки по координатам из списка.
            """
            unique_points = []
            seen_coords = set()
            for point in points_list:
                coord_key = (round(point.latitude, 6), round(point.longitude, 6))
                if coord_key not in seen_coords:
                    seen_coords.add(coord_key)
                    unique_points.append(point)
            return unique_points

        # --- Формирование отчета по регионам России ---
        for region in sorted(russia_by_region.keys()):
            if russia_by_region[region]:
                # Добавляем заголовок региона
                if region.startswith("на территории"):
                    add_region_header(f"{region}:")
                else:
                    add_region_header(f"на территории {region}:")
                unique_points = get_unique_points(russia_by_region[region])
                for point in unique_points:
                    if point.area_desc:
                        add_normal_paragraph(f"{point.area_desc};")

        # --- Формирование отчета по точкам Украины ---
        if ukraine_points:
            add_region_header("на территории Украины:")
            unique_ukraine_points = get_unique_points(ukraine_points)
            for point in unique_ukraine_points:
                if point.area_desc:
                    add_normal_paragraph(f"{point.area_desc};")

        # --- Формирование отчета по точкам других стран ---
        if other_points:
            other_by_country: Dict[str, List[PointRecord]] = {}
            for point in other_points:
                country = point.country or "Неизвестная страна"
                if country not in other_by_country:
                    other_by_country[country] = []
                other_by_country[country].append(point)
            from src.core import get_country_by_lat_lon

            for country, points in sorted(other_by_country.items()):
                if points:
                    # Получаем русское название страны по координатам первой точки
                    country_rus = get_country_by_lat_lon(points[0].latitude, points[0].longitude)[1]
                    add_region_header(f"на территории {country_rus}:")
                    for point in get_unique_points(points):
                        if point.area_desc:
                            add_normal_paragraph(f"{point.area_desc};")

        # --- Добавляем информацию о не найденных городах ---
        if wrong_city_data_folder:
            doc.add_paragraph("")
            doc.add_paragraph(
                "Города не найдены в city.txt. Требуется описание и повторный запуск обработки файлов директории (перед повторным запуском удалите data.xlsx в директории):"
            )
            for city in wrong_city_data_folder:
                add_normal_paragraph(city)
            doc.add_paragraph("")

        # --- Добавляем оригинальные тексты точек ---
        if all_original_texts:
            doc.add_paragraph("")
            doc.add_paragraph("")
            for i, original_text in enumerate(all_original_texts):
                if i > 0:
                    doc.add_paragraph("")
                formatted_text = format_original_text(original_text)
                if formatted_text:
                    lines = formatted_text.split("\n")
                    for line in lines:
                        if line.strip():
                            add_normal_paragraph(line)

        # --- Сохраняем документ ---
        os.makedirs(os.path.dirname(report_path), exist_ok=True)
        doc.save(report_path)
        if log_message:
            log_message(
                f"Word отчет успешно создан: {report_path}", logger_level="info", color="blue"
            )
        return True

    except ImportError as e:
        # Обработка ошибки отсутствия модуля python-docx
        if log_message:
            log_message(
                f"Модуль python-docx не установлен: {e}", logger_level="error", color="red"
            )
            log_message(
                "Установите модуль командой: pip install python-docx",
                logger_level="info",
                color="blue",
            )
        return False
    except Exception as e:
        # Обработка других ошибок при создании отчета
        if log_message:
            log_message(
                f"Ошибка при создании Word отчета {report_path}: {e}",
                logger_level="error",
                color="red",
            )
        return False
