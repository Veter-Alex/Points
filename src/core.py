import csv
import os
import shutil
import urllib.request
import xml.etree.ElementTree as ET
from typing import Dict, List, Optional

import geopandas as gpd
from shapely.geometry import Point

from models.city import CityData, CityRecord
from models.points import PointRecord, PointsData
from src.logger import logger

_WORLD_GDF = None  # Кэш для GeoDataFrame


def find_folders_missing_data_csv(rootFolder: str) -> List[str]:
    """
    Ищет во всех поддиректориях rootFolder папки, в которых есть хотя бы один xml или json файл,
    но нет файла data.xlsx. Папки с именем 'bad' игнорируются.
    Возвращает список абсолютных путей к таким папкам.
    """
    result = []
    for dirpath, dirnames, filenames in os.walk(rootFolder):
        # Пропустить саму папку bad (только если текущая папка называется bad)
        if os.path.basename(dirpath).lower() == "bad":
            continue
        has_xml_or_json = any(f.lower().endswith((".xml", ".json")) for f in filenames)
        has_data_csv = any(f.lower() == "data.xlsx" for f in filenames)
        if has_xml_or_json and not has_data_csv:
            result.append(dirpath)
    return result


def parse_xml(xml_path: str) -> Optional[PointRecord]:
    """
    Универсальный парсер xml-файла с точкой.
    Возвращает PointRecord (или None, если невалидно).
    """
    logger.info(f"Начинаем обработку XML файла: {xml_path}")
    try:
        # получаем текст XML
        import json

        with open(xml_path, encoding="utf-8") as f_xml:
            xml_text = f_xml.read()

        # парсим XML
        tree = ET.ElementTree(ET.fromstring(xml_text))
        root = tree.getroot()
        # --- Формат 1: <root><point>...
        point = root.find(".//point")
        if point is not None:
            lat = point.find("latitude")
            lat_val = lat.attrib.get("Value") if lat is not None else None
            lon = point.find("longitude")
            lon_val = lon.attrib.get("Value") if lon is not None else None
            dt = point.find("datetime")
            dt_val = dt.attrib.get("Value") if dt is not None else None
            city = point.find("City")
            if city is not None:
                city_val = city.attrib.get("Value")
            else:
                name = point.find("name")
                city_val = name.attrib.get("Value") if name is not None else None
            country = point.find("Country")
            country_val = country.attrib.get("Value") if country is not None else None
            # Разделить дату и время
            date_val, time_val = None, None
            if dt_val and " " in dt_val:
                date_val, time_val = dt_val.split(" ", 1)
            elif dt_val:
                date_val = dt_val
            # Проверка обязательных
            if not (lat_val and lon_val and date_val and time_val):
                raise ValueError("Нет обязательных данных (lat/lon/date/time)")
            return PointRecord(
                date=date_val,
                time=time_val,
                latitude=float(lat_val),
                longitude=float(lon_val),
                x_sk42=None,
                y_sk42=None,
                country=country_val,
                city=city_val,
                area_desc=None,
                region_desc=None,
                original_text=xml_text,
                file_path=xml_path,
            )
        # --- Формат 2: <weather><loc ...><obs .../></loc></weather>
        if root.tag == "weather":
            loc = root.find(".//loc")
            if loc is None:
                raise ValueError("Нет тега <loc>")
            city_val = loc.attrib.get("name")
            country_val = loc.attrib.get("country")
            lat_val = loc.attrib.get("lat")
            lon_val = loc.attrib.get("lon")
            # Найти первую подходящую запись с датой и временем
            dt_val = None
            for tag in ["obs", "latest", "fcd", "fch"]:
                node = loc.find(tag)
                if node is not None and "dt" in node.attrib:
                    dt_val = node.attrib["dt"]
                    break
            if not dt_val:
                raise ValueError("Нет даты и времени (dt)")
            # Разделить дату и время
            date_val, time_val = None, None
            if "T" in dt_val:
                date_val, time_val = dt_val.split("T", 1)
                time_val = time_val.split("Z")[0]  # убрать Z, если есть
            elif " " in dt_val:
                date_val, time_val = dt_val.split(" ", 1)
            else:
                date_val = dt_val
            if not (lat_val and lon_val and date_val and time_val):
                raise ValueError("Нет обязательных данных (lat/lon/date/time)")
            return PointRecord(
                date=date_val,
                time=time_val,
                latitude=float(lat_val),
                longitude=float(lon_val),
                x_sk42=None,
                y_sk42=None,
                country=country_val,
                city=city_val,
                area_desc=None,
                region_desc=None,
                original_text=xml_text,
                file_path=xml_path,
            )
        # --- Формат 3: <document><GetHHForecastResult ...>
        if root.tag == "document":
            hh = root.find(".//GetHHForecastResult")
            if hh is not None:
                city_val = hh.attrib.get("cityName")
                country_val = hh.attrib.get("country_name")
                lat_val = hh.attrib.get("lat")
                lon_val = hh.attrib.get("lng")
                # Найти первую дату/время: <forecastData><HHWeather><time>...</time>
                dt_val = None
                # 1. <forecastData><HHWeather><time>
                hhweather = hh.find("forecastData/HHWeather/time")
                if hhweather is not None and hhweather.text:
                    dt_val = hhweather.text.strip()
                # 2. <forecastData><HHForecasts><HHForecast><time>
                if not dt_val:
                    hhforecast = hh.find("forecastData/HHForecasts/HHForecast/time")
                    if hhforecast is not None and hhforecast.text:
                        dt_val = hhforecast.text.strip()
                # Разделить дату и время
                date_val, time_val = None, None
                if dt_val and " " in dt_val:
                    date_val, time_val = dt_val.split(" ", 1)
                elif dt_val:
                    date_val = dt_val
                if not (lat_val or lon_val or date_val or time_val):
                    raise ValueError("Нет обязательных данных (lat/lon/date/time)")
            return PointRecord(
                date=date_val,
                time=time_val,
                latitude=float(lat_val),
                longitude=float(lon_val),
                x_sk42=None,
                y_sk42=None,
                country=country_val,
                city=city_val,
                area_desc=None,
                region_desc=None,
                original_text=xml_text,
                file_path=xml_path,
            )
        # --- Неизвестный формат
        raise ValueError("Неизвестный формат XML")
    except Exception as e:
        # Переместить файл в bad
        bad_dir = os.path.join(os.path.dirname(xml_path), "bad")
        os.makedirs(bad_dir, exist_ok=True)
        shutil.move(xml_path, os.path.join(bad_dir, os.path.basename(xml_path)))
        # Переместить одноименный .spr файл, если он существует
        spr_path = os.path.splitext(xml_path)[0] + ".spr"
        if os.path.exists(spr_path):
            shutil.move(spr_path, os.path.join(bad_dir, os.path.basename(spr_path)))
        return None
    finally:
        logger.info(f"Файл успешно обработан: {xml_path}")


def find_point_by_lat_lon(
    points_data: List[PointRecord], latitude: float, longitude: float, tol: float = 1e-6
):
    """
    Найти точку в points_data по координатам (с учетом допуска tol).
    Возвращает PointRecord или None.
    """
    for point in points_data:
        # point.latitude, point.longitude должны быть float
        if (
            abs(point.latitude - latitude) <= tol
            and abs(point.longitude - longitude) <= tol
        ):
            return point
    return None


from typing import Tuple


def get_country_by_lat_lon(lat: float, lon: float) -> Tuple[str, str]:
    """
    Определить страну по координатам (lat, lon) с помощью границ стран (GeoJSON).
    При первом вызове скачивает и кэширует countries.geojson.
    Возвращает название страны или пустую строку.
    """
    global _WORLD_GDF
    if _WORLD_GDF is None:
        geojson_path = os.path.join(os.path.dirname(__file__), "countries.geojson")
        _WORLD_GDF = gpd.read_file(geojson_path)

    point = Point(lon, lat)
    matches = _WORLD_GDF[_WORLD_GDF["geometry"].contains(point)]
    country_eng = ""
    country_rus = ""

    if not matches.empty:
        # Используем колонку 'name' из countries.geojson
        country_eng = matches.iloc[0]["name"]
    else:
        # Если точное попадание не найдено, ищем в небольшом радиусе (для точек на границах)
        buffer_point = point.buffer(0.02)  # буфер 0.02 градуса (~2.2 км)
        buffer_matches = _WORLD_GDF[_WORLD_GDF["geometry"].intersects(buffer_point)]
        if not buffer_matches.empty:
            country_eng = buffer_matches.iloc[0]["name"]

    # Словарь переводов: Английское название -> Русское в предложном падеже
    country_translate = {
        "Russia": "России",
        "Ukraine": "Украины",
        "Germany": "Германии",
        "France": "Франции",
        "China": "Китая",
        "Belarus": "Белоруссии",
        "Poland": "Польши",
        "Kazakhstan": "Казахстана",
        "United States of America": "США",
        "United States": "США",
        "USA": "США",
        "Turkey": "Турции",
        "Italy": "Италии",
        "Spain": "Испании",
        "United Kingdom": "Великобритании",
        "Czechia": "Чехии",
        "Czech Republic": "Чехии",
        "Estonia": "Эстонии",
        "Latvia": "Латвии",
        "Lithuania": "Литвы",
        "Finland": "Финляндии",
        "Georgia": "Грузии",
        "Armenia": "Армении",
        "Azerbaijan": "Азербайджана",
        "Moldova": "Молдовы",
        "Uzbekistan": "Узбекистана",
        "Kyrgyzstan": "Киргизии",
        "Tajikistan": "Таджикистана",
        "Turkmenistan": "Туркменистана",
        "Romania": "Румынии",
        "Bulgaria": "Болгарии",
        "Serbia": "Сербии",
        "Hungary": "Венгрии",
        "Slovakia": "Словакии",
        "Slovenia": "Словении",
        "Croatia": "Хорватии",
        "Montenegro": "Черногории",
        "Bosnia and Herzegovina": "Боснии и Герцеговины",
        "North Macedonia": "Северной Македонии",
        "Greece": "Греции",
        "Sweden": "Швеции",
        "Norway": "Норвегии",
        "Denmark": "Дании",
        "Iceland": "Исландии",
        "Malta": "Мальты",
        "Liechtenstein": "Лихтенштейна",
        "Monaco": "Монако",
        "San Marino": "Сан-Марино",
        "Vatican City": "Ватикана",
        "Switzerland": "Швейцарии",
        "Luxembourg": "Люксембурга",
        "Singapore": "Сингапура",
        "Japan": "Японии",
        "South Korea": "Южной Кореи",
        "North Korea": "Северной Кореи",
        "Thailand": "Таиланда",
        "Vietnam": "Вьетнама",
        "Malaysia": "Малайзии",
        "Indonesia": "Индонезии",
        "Philippines": "Филиппин",
        "India": "Индии",
        "Pakistan": "Пакистана",
        "Bangladesh": "Бангладеш",
        "Sri Lanka": "Шри-Ланки",
        "Australia": "Австралии",
        "New Zealand": "Новой Зеландии",
        "Canada": "Канады",
        "Mexico": "Мексики",
        "Brazil": "Бразилии",
        "Argentina": "Аргентины",
        "Chile": "Чили",
        "Peru": "Перу",
        "Colombia": "Колумбии",
        "Venezuela": "Венесуэлы",
        "Ecuador": "Эквадора",
        "Bolivia": "Боливии",
        "Paraguay": "Парагвая",
        "Uruguay": "Уругвая",
        "Egypt": "Египта",
        "South Africa": "Южной Африки",
        "Morocco": "Марокко",
        "Algeria": "Алжира",
        "Tunisia": "Туниса",
        "Libya": "Ливии",
        "Sudan": "Судана",
        "Ethiopia": "Эфиопии",
        "Kenya": "Кении",
        "Tanzania": "Танзании",
        "Nigeria": "Нигерии",
        "Ghana": "Ганы",
        "Iran": "Ирана",
        "Iraq": "Ирака",
        "Syria": "Сирии",
        "Lebanon": "Ливана",
        "Jordan": "Иордании",
        "Israel": "Израиля",
        "Saudi Arabia": "Саудовской Аравии",
        "UAE": "ОАЭ",
        "Qatar": "Катара",
        "Kuwait": "Кувейта",
        "Bahrain": "Бахрейна",
        "Oman": "Омана",
        "Yemen": "Йемена",
        "Afghanistan": "Афганистана",
        "Nepal": "Непала",
        "Bhutan": "Бутана",
        "Myanmar": "Мьянмы",
        "Cambodia": "Камбоджи",
        "Laos": "Лаоса",
        # ... можно добавить другие страны по необходимости ...
    }
    country_rus = country_translate.get(country_eng, country_eng)
    return country_eng, country_rus


def get_sk42_coordinates(lat: float, lon: float):
    """Преобразование координат WGS84 -> СК-42 с автоопределением зоны."""
    import math

    # Проверяем, что координаты в допустимом диапазоне для СК-42
    # СК-42 покрывает территорию России от 18° до 165° восточной долготы
    if lon < 18 or lon > 165:
        logger.warning(f"Координаты вне зоны действия СК-42: {lat}, {lon}")
        x_sk42, y_sk42 = None, None
    else:
        try:
            # Преобразование координат WGS84 -> СК-42
            from src.coordinate_transformer import CoordinateTransformer

            transformer = CoordinateTransformer(system="SK42_GAUSS_KRUGER", zone="AUTO")
            x_sk42, y_sk42 = transformer.transform(lat, lon, to_wgs=False)
            # Проверка на бесконечность и NaN
            if x_sk42 is not None and (math.isinf(x_sk42) or math.isnan(x_sk42)):
                x_sk42 = None
            if y_sk42 is not None and (math.isinf(y_sk42) or math.isnan(y_sk42)):
                y_sk42 = None
        except Exception as e:
            logger.error(f"Ошибка преобразования координат WGS84->СК-42: {e}")
            x_sk42, y_sk42 = None, None
        x_sk42 = int(x_sk42) if x_sk42 is not None else None
        y_sk42 = int(y_sk42) if y_sk42 is not None else None
    # Возвращаем преобразованные координаты
    return x_sk42, y_sk42


def find_city_by_name(records, city_name):
    """
    Найти город в records по имени.
    Возвращает CityRecord или None.
    """
    if not city_name:
        return None
    for record in records:
        name_ru = record.name_ru if record.name_ru else ""
        name_original = record.name_original if record.name_original else ""
        if (
            name_ru.lower() == city_name.lower()
            or name_original.lower() == city_name.lower()
        ):
            return record
    return None


def get_area_desc(
    found_city: CityRecord, result_parse_xml: PointRecord, x_sk42: int, y_sk42: int
) -> str:
    """
    Сформировать описание района (area_desc) для точки из city_data:
    район 6 км северо-западнее н.п.Суджа (87 км юго-зап. г.Курск, координаты: X=5681162 Y=6655966);
    Sudzha=н.п.Суджа_51.190581_35.269206_Россия_87 км юго-зап. г.Курск_на территории Курской области
    Использует данные из found_city и result_parse_xml для создания описания.
    """
    import math

    def haversine(lat1, lon1, lat2, lon2):
        # расстояние между двумя точками на сфере (км)
        R = 6371.0
        phi1 = math.radians(lat1)
        phi2 = math.radians(lat2)
        dphi = math.radians(lat2 - lat1)
        dlambda = math.radians(lon2 - lon1)
        a = (
            math.sin(dphi / 2) ** 2
            + math.cos(phi1) * math.cos(phi2) * math.sin(dlambda / 2) ** 2
        )
        c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))
        return R * c

    def get_direction(lat1, lon1, lat2, lon2):
        # возвращает направление (например, "северо-западнее")
        dlat = lat2 - lat1
        dlon = lon2 - lon1
        angle = math.degrees(math.atan2(dlon, dlat))
        dirs = [
            (22.5, "севернее"),
            (67.5, "северо-восточнее"),
            (112.5, "восточнее"),
            (157.5, "юго-восточнее"),
            (202.5, "южнее"),
            (247.5, "юго-западнее"),
            (292.5, "западнее"),
            (337.5, "северо-западнее"),
            (360.0, "севернее"),
        ]
        angle = (angle + 360) % 360
        for bound, name in dirs:
            if angle < bound:
                return name
        return "севернее"

    # Получить расстояние от точки из result_parse_xml до найденного города
    distance = haversine(
        found_city.latitude,
        found_city.longitude,
        result_parse_xml.latitude,
        result_parse_xml.longitude,
    )
    # Получить направление от города до точки
    direction = get_direction(
        found_city.latitude,
        found_city.longitude,
        result_parse_xml.latitude,
        result_parse_xml.longitude,
    )

    # Если расстояние меньше 2 км
    if distance <= 2:
        # Если найденный город не имеет описание района
        if found_city.description is None:
            # вернуть описание в следующем формате
            # район г.Сумы (координаты: X=5645005 Y=6626587);
            return f"район {found_city.name_ru} (координаты: X={x_sk42} Y={y_sk42})"
        # Если найденный город имеет описание района
        else:
            # вернуть описание в следующем формате
            # район г.Сумы (87 км юго-зап. г.Курск, координаты: X=5673347 Y=6663897);
            return f"район {found_city.name_ru} ({found_city.description}, координаты: X={x_sk42} Y={y_sk42})"

    # если расстояние больше 2 км
    elif distance > 2:
        # Если найденный город не имеет описание района
        if found_city.description is None:
            # вернуть описание в следующем формате
            # район 5,8 км юго-западнее н.п.Суджа (координаты: X=5673347 Y=6663897);
            return f"район {distance:.1f} км {direction} {found_city.name_ru} (координаты: X={x_sk42} Y={y_sk42})"
        # Если найденный город имеет описание района
        else:
            # вернуть описание в следующем формате
            # район 5,8 км юго-западнее н.п.Суджа (87 км юго-зап. г.Курск, координаты: X=5673347 Y=6663897);
            return f"район {distance:.1f} км {direction} {found_city.name_ru} ({found_city.description}, координаты: X={x_sk42} Y={y_sk42})"
    # На случай, если ни одно условие не сработало, возвращаем пустую строку
    return "error in get_area_desc"


def new_point_from_city_data(
    found_city: CityRecord, result_parse_xml: PointRecord
) -> PointRecord:
    """
    Создать новый PointRecord из результата парсинга XML.
    Использует данные из result_parse_xml и city_data для создания нового объекта.
    """
    country_eng, country_rus = get_country_by_lat_lon(
        result_parse_xml.latitude, result_parse_xml.longitude
    )

    x_sk42, y_sk42 = get_sk42_coordinates(
        result_parse_xml.latitude, result_parse_xml.longitude
    )

    return PointRecord(
        date=result_parse_xml.date,
        time=result_parse_xml.time,
        latitude=result_parse_xml.latitude,
        longitude=result_parse_xml.longitude,
        x_sk42=x_sk42,
        y_sk42=y_sk42,
        country=country_eng,
        city=result_parse_xml.city,
        area_desc=get_area_desc(found_city, result_parse_xml, x_sk42, y_sk42),
        region_desc=found_city.region if found_city else None,
        original_text=result_parse_xml.original_text,
        file_path=result_parse_xml.file_path,
    )


def new_point_from_points_data(
    found_point: PointRecord, result_parse_xml: PointRecord
) -> PointRecord:
    """
    Создать новый PointRecord из найденной точки и результата парсинга XML.
    Использует данные из found_point и result_parse_xml для создания нового объекта.
    """
    return PointRecord(
        date=result_parse_xml.date,
        time=result_parse_xml.time,
        latitude=found_point.latitude,
        longitude=found_point.longitude,
        x_sk42=found_point.x_sk42,
        y_sk42=found_point.y_sk42,
        country=found_point.country,
        city=found_point.city,
        area_desc=found_point.area_desc,
        region_desc=found_point.region_desc,
        original_text=result_parse_xml.original_text,
        file_path=result_parse_xml.file_path,
    )

def new_point_without_city(result_parse_xml: PointRecord) -> PointRecord:
    """
    Создать новый PointRecord, если в результатах парсинга город не найден и
    нет информации с координатами в ранее отмеченных точках (points_data).
    """
    # TODO проверить логику работы функции

    country_eng, country_rus = get_country_by_lat_lon(
        result_parse_xml.latitude, result_parse_xml.longitude
    )
    x_sk42, y_sk42 = get_sk42_coordinates(
        result_parse_xml.latitude, result_parse_xml.longitude
    )
    # return PointRecord(
    #     date=result_parse_xml.date,
    #     time=result_parse_xml.time,
    #     latitude=result_parse_xml.latitude,
    #     longitude=result_parse_xml.longitude,
    #     x_sk42=x_sk42,
    #     y_sk42=y_sk42,
    #     country=country_eng,
    #     city=None,
    #     area_desc=None,
    #     region_desc=None,
    #     original_text=result_parse_xml.original_text,
    #     file_path=result_parse_xml.file_path,
    # )
    return None

def save_points_to_excel(points_folder: List[PointRecord], data_xlsx_path: str) -> bool:
    """
    Сохраняет список PointRecord в Excel с форматированием переносов строк.
    
    Args:
        points_folder: Список объектов PointRecord для сохранения
        data_xlsx_path: Путь к файлу Excel для сохранения
        
    Returns:
        bool: True если сохранение успешно, False если произошла ошибка
    """
    import pandas as pd
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
    
    try:
        # Подготовка данных с нормализацией переносов строк
        df_data = []
        for point in points_folder:
            row_data = point.__dict__.copy()
            if row_data.get("original_text"):
                # Нормализация переносов строк: все виды -> \n
                text = str(row_data["original_text"])
                row_data["original_text"] = text.replace("\r\n", "\n").replace("\r", "\n")
            df_data.append(row_data)
        
        points_df = pd.DataFrame(df_data)
        
        # Сохранение с openpyxl для поддержки форматирования
        with pd.ExcelWriter(data_xlsx_path, engine='openpyxl') as writer:
            points_df.to_excel(writer, index=False, sheet_name='Points')
            worksheet = writer.sheets['Points']
            
            # Стили для таблицы
            header_font = Font(bold=True, name='Times New Roman', size=11)
            data_font = Font(name='Times New Roman', size=10)
            header_fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")
            center_alignment = Alignment(horizontal='center', vertical='center')
            wrap_alignment = Alignment(wrapText=True, vertical='center', horizontal='center')
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # Форматирование заголовков (первая строка)
            for col_num in range(1, len(points_df.columns) + 1):
                cell = worksheet.cell(row=1, column=col_num)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = center_alignment
                cell.border = thin_border
            
            # Форматирование данных и границы для всех ячеек
            for row_num in range(2, worksheet.max_row + 1):
                for col_num in range(1, len(points_df.columns) + 1):
                    cell = worksheet.cell(row=row_num, column=col_num)
                    cell.border = thin_border
                    cell.alignment = center_alignment
                    cell.font = data_font
            
            # Специальное форматирование для столбца original_text
            original_text_col = None
            if 'original_text' in points_df.columns:
                original_text_col = points_df.columns.get_loc('original_text') + 1
                
                # Применяем перенос текста для столбца original_text
                for row_num in range(2, worksheet.max_row + 1):
                    cell = worksheet.cell(row=row_num, column=original_text_col)
                    if cell.value:
                        cell.alignment = wrap_alignment
                        cell.font = data_font
                
                # Установка фиксированной ширины для столбца original_text
                col_letter = worksheet.cell(row=1, column=original_text_col).column_letter
                worksheet.column_dimensions[col_letter].width = 50
            
            # Автоподбор ширины для всех остальных столбцов (альтернативный метод)
            for col_num in range(1, len(points_df.columns) + 1):
                if original_text_col is None or col_num != original_text_col:
                    col_letter = worksheet.cell(row=1, column=col_num).column_letter
                    
                    # Простой автоподбор: находим максимальную длину в столбце
                    max_length = 0
                    for row in worksheet[col_letter]:
                        try:
                            if len(str(row.value)) > max_length:
                                max_length = len(str(row.value))
                        except:
                            pass
                    
                    # Устанавливаем ширину с небольшим запасом (минимум 8, максимум 50)
                    adjusted_width = min(max(max_length + 2, 8), 50)
                    worksheet.column_dimensions[col_letter].width = adjusted_width
        
        logger.info(f"Excel файл успешно сохранён: {data_xlsx_path}")
        return True
        
    except Exception as e:
        logger.warning(f"Ошибка при сохранении Excel: {e}")
        return False


def create_kml_file(point: PointRecord, kml_file_path: str) -> bool:
    """
    Создает KML файл для точки с иерархической структурой папок по дате и времени.
    
    Args:
        point: PointRecord - объект точки для создания KML
        kml_file_path: str - путь к файлу KML для сохранения
        
    Returns:
        bool: True если создание успешно, False если произошла ошибка
    """
    try:
        # Извлекаем компоненты даты и времени
        date_parts = point.date.split(".")
        time_parts = point.time.split(":")
        
        if len(date_parts) >= 3 and len(time_parts) >= 2:
            day = date_parts[0]
            month = date_parts[1] 
            year = date_parts[2]
            hour = time_parts[0]
            minute = time_parts[1]
        else:
            logger.warning(f"Неверный формат даты/времени: {point.date} {point.time}")
            return False
        
        # Формируем содержимое KML файла
        kml_content = f'''<?xml version="1.0" encoding="UTF-8"?>
<kml xmlns="http://earth.google.com/kml/2.2">
  <Document>
    <Folder>
      <name>{year}</name>
      <open>1</open>
      <Style>
        <ListStyle>
          <listItemType>check</listItemType>
          <bgColor>00ffffff</bgColor>
        </ListStyle>
      </Style>
      <Folder>
        <name>{month}</name>
        <open>1</open>
        <Style>
          <ListStyle>
            <listItemType>check</listItemType>
            <bgColor>00ffffff</bgColor>
          </ListStyle>
        </Style>
        <Folder>
          <name>{day}</name>
          <open>1</open>
          <Style>
            <ListStyle>
              <listItemType>check</listItemType>
              <bgColor>00ffffff</bgColor>
            </ListStyle>
          </Style>
          <Placemark>
            <name>{point.date} {point.time} (X={point.x_sk42 or 'N/A'} Y={point.y_sk42 or 'N/A'})
</name>
            <description>{point.date} {point.time}
City binding - {point.city or 'Unknown'}
Country - {point.country or 'Unknown'}
Latitude_SK42_GEO - {point.latitude:.4f}
Longitude_SK42_GEO - {point.longitude:.4f}
Latitude_SK42_Gauss_Kruger - {point.x_sk42 or 'N/A'}
Longitude_SK42_Gauss_Kruger - {point.y_sk42 or 'N/A'}</description>
            <Style>
              <LabelStyle>
                <color>FF00FFFF</color>
                <scale>1.09090909090909</scale>
              </LabelStyle>
              <IconStyle>
                <scale>0.390625</scale>
                <Icon>
                  <href>files/1.png</href>
                </Icon>
                <hotSpot x="0.5" y="0" xunits="fraction" yunits="fraction"/>
              </IconStyle>
            </Style>
            <Point>
              <extrude>1</extrude>
              <coordinates>{point.longitude:.8f},{point.latitude:.8f},0</coordinates>
            </Point>
          </Placemark>
        </Folder>
      </Folder>
    </Folder>
  </Document>
</kml>'''
        
        # Создаем директорию, если не существует
        os.makedirs(os.path.dirname(kml_file_path), exist_ok=True)
        
        # Записываем KML файл
        with open(kml_file_path, 'w', encoding='utf-8') as f:
            f.write(kml_content)
        
        logger.info(f"KML файл успешно создан: {kml_file_path}")
        return True
        
    except Exception as e:
        logger.error(f"Ошибка при создании KML файла {kml_file_path}: {e}")
        return False


def create_word_report(points_folder: List[PointRecord], report_path: str) -> bool:
    """
    Создает Word отчет с точками, сгруппированными по странам и областям.
    Сначала точки на территории России (разные области), потом Украины, потом остальные страны.
    
    Args:
        points_folder: Список объектов PointRecord для включения в отчет
        report_path: Путь к файлу Word для сохранения
        
    Returns:
        bool: True если создание успешно, False если произошла ошибка
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        
        # Создаем новый документ
        doc = Document()
        
        # Настройка стандартного стиля документа
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Группируем точки по странам и областям
        russia_points = []
        ukraine_points = []
        other_points = []
        all_original_texts = []  # Для сбора всех original_text
        seen_coordinates = set()  # Для отслеживания уникальных координат
        
        for point in points_folder:
            # Создаем ключ для координат (округляем до 6 знаков после запятой)
            coord_key = (round(point.latitude, 6), round(point.longitude, 6))
            
            # Добавляем original_text только от уникальных координат
            if coord_key not in seen_coordinates:
                seen_coordinates.add(coord_key)
                if point.original_text:
                    all_original_texts.append(point.original_text)
            
            if point.country == "Russia":
                russia_points.append(point)
            elif point.country == "Ukraine":
                ukraine_points.append(point)
            else:
                other_points.append(point)
        
        # Функция для извлечения области из area_desc
        def extract_region(area_desc: str) -> str:
            if not area_desc:
                return "Неизвестная область"
            
            # Ищем паттерн "на территории [название области]"
            import re
            pattern = r'на территории (.+?)(?:,|$)'
            match = re.search(pattern, area_desc)
            if match:
                return match.group(1).strip()
            
            # Если не найден паттерн, ищем область в region_desc
            return "Неизвестная область"
        
        # Функция для создания абзаца с курсивом для заголовков
        def add_region_header(text: str):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.italic = True
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return p
        
        # Функция для создания обычного абзаца
        def add_normal_paragraph(text: str):
            p = doc.add_paragraph(text)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            return p
        
        # Функция для форматирования original_text
        def format_original_text(text: str) -> str:
            if not text:
                return ""
            
            # Разбиваем на строки и обрабатываем каждую
            lines = text.split('\n')
            formatted_lines = []
            
            for line in lines:
                # Удаляем пробелы в начале и конце строки
                clean_line = line.strip()
                if clean_line:  # Добавляем только непустые строки
                    formatted_lines.append(clean_line)
            
            # Соединяем строки с обычным переводом строки
            return '\n'.join(formatted_lines)
        
        # Группируем российские точки по областям
        russia_by_region = {}
        for point in russia_points:
            region = point.region_desc if point.region_desc else extract_region(point.area_desc or "")
            if region not in russia_by_region:
                russia_by_region[region] = []
            russia_by_region[region].append(point)
        
        # Функция для получения уникальных точек по координатам
        def get_unique_points(points_list):
            unique_points = []
            seen_coords = set()
            for point in points_list:
                coord_key = (round(point.latitude, 6), round(point.longitude, 6))
                if coord_key not in seen_coords:
                    seen_coords.add(coord_key)
                    unique_points.append(point)
            return unique_points
        
        # Добавляем российские точки по областям
        for region in sorted(russia_by_region.keys()):
            if russia_by_region[region]:
                # Заголовок области (курсив) - убираем дублирование "на территории"
                if region.startswith("на территории"):
                    add_region_header(f"{region}:")
                else:
                    add_region_header(f"на территории {region}:")
                
                # Получаем уникальные точки в области
                unique_points = get_unique_points(russia_by_region[region])
                for point in unique_points:
                    if point.area_desc:
                        add_normal_paragraph(f"{point.area_desc};")
        
        # Добавляем украинские точки
        if ukraine_points:
            # Заголовок для Украины (курсив)
            add_region_header("на территории Украины:")
            
            # Получаем уникальные украинские точки
            unique_ukraine_points = get_unique_points(ukraine_points)
            for point in unique_ukraine_points:
                if point.area_desc:
                    add_normal_paragraph(f"{point.area_desc};")
        
        # Добавляем точки других стран
        if other_points:
            # Группируем по странам
            other_by_country = {}
            for point in other_points:
                country = point.country or "Неизвестная страна"
                if country not in other_by_country:
                    other_by_country[country] = []
                other_by_country[country].append(point)
            
            # Добавляем по странам
            for country in sorted(other_by_country.keys()):
                if other_by_country[country]:
                    # Заголовок страны (курсив)
                    country_rus = get_country_by_lat_lon(
                        other_by_country[country][0].latitude, 
                        other_by_country[country][0].longitude
                    )[1]
                    add_region_header(f"на территории {country_rus}:")
                    
                    # Получаем уникальные точки в стране
                    unique_country_points = get_unique_points(other_by_country[country])
                    for point in unique_country_points:
                        if point.area_desc:
                            add_normal_paragraph(f"{point.area_desc};")
        
        # Добавляем 2 пустые строки и original_text для всех точек
        if all_original_texts:
            # Добавляем 2 пустые строки
            doc.add_paragraph("")
            doc.add_paragraph("")
            
            # Добавляем каждый original_text через одну пустую строку
            for i, original_text in enumerate(all_original_texts):
                if i > 0:  # Добавляем пустую строку между текстами (кроме первого)
                    doc.add_paragraph("")
                
                # Форматируем original_text и добавляем как отдельные строки
                formatted_text = format_original_text(original_text)
                if formatted_text:
                    # Разбиваем на строки и добавляем каждую как отдельный абзац
                    lines = formatted_text.split('\n')
                    for line in lines:
                        if line.strip():  # Добавляем только непустые строки
                            add_normal_paragraph(line)
        
        # Создаем директорию, если не существует
        os.makedirs(os.path.dirname(report_path), exist_ok=True)
        
        # Сохраняем документ
        doc.save(report_path)
        
        logger.info(f"Word отчет успешно создан: {report_path}")
        return True
        
    except ImportError as e:
        logger.error(f"Модуль python-docx не установлен: {e}")
        logger.info("Установите модуль командой: pip install python-docx")
        return False
    except Exception as e:
        logger.error(f"Ошибка при создании Word отчета {report_path}: {e}")
        return False
        
        
def find_and_parse_files(
    input_folder: str, city_data: CityData, points_data: PointsData
) -> None:
    """
    Основная функция для поиска папок с отсутствующим data.xlsx и парсинга xml и json файлов.

    input_folder: str - путь к корневой папке для поиска
    city_data: CityData - объект с данными о городах
    points_data: PointsData - объект с данными о всех ранее отмеченных точках
    """
    # Список для городов, которые не удалось найти в city_data
    wrong_city_data = []

    # Запустить поиск папок с отсутствующим data.xlsx
    for folder in find_folders_missing_data_csv(input_folder):
        print(f"Папка без data.xlsx: {folder}")

        # список для сбора точек в текущей папке
        points_folder: list[PointRecord] = []

        # Пройтись по всем файлам в папке
        for file in os.listdir(folder):
            new_point = None
            result_parse = None  # результат парсинга файла

            if file.endswith(".xml"):
                # Парсинг XML файла
                result_parse = parse_xml(os.path.join(folder, file))

            if file.endswith(".json"):
                # Парсинг JSON файла
                pass
                # result_parse = parse_json_file(os.path.join(folder, file))

            if result_parse is not None:
            # если парсинг успешен (есть координаты, дата и время)
                if result_parse.city is not None:
                # если в результате парсинга есть название города

                    # найти название города в city_data
                    found_city = find_city_by_name(
                        city_data.records, result_parse.city
                    )
                    if found_city:
                        # Если город найден в city_data
                        # создаём новую точку на основе найденной записи
                        new_point = new_point_from_city_data(
                            found_city, result_parse
                        )

                    else:
                        # Если город не найден (но есть название)
                        # создаём новую запись о городе вида
                        # Druzhba=н.п.Дружба_52.041342_33.943643_Украина_140 км сев.-зап. г.Сумы_на территории Украины
                        # и добавляем в wrong_city_data для дальнейшего анализа
                        logger.warning(
                            f"Город не найден в city_data: {result_parse.city}"
                        )
                        # получить название страны по координатам
                        country_eng, country_rus = get_country_by_lat_lon(
                            result_parse.latitude, result_parse.longitude
                        )
                        # создать новую запись города
                        new_city = (
                            f"{result_parse.city}=н.п.НАЗВАНИЕ ГОРОДА_"
                            f"широта центра города_долгота центра города_"
                            f"{country_eng}_описание относитьно обласного центра_"
                            f"на территории {country_rus}"
                        )
                        # добавить в список таких городов
                        wrong_city_data.append(new_city)
                        logger.info(f"Добавлен город для анализа: {new_city}")
                
                else:
                    # Если город не указан в точке, но есть координаты
    
                    # Ищем точку по координатам ранее отмеченных точках
                    found_point = find_point_by_lat_lon(
                        points_data.points,
                        result_parse.latitude,
                        result_parse.longitude,
                    )
                    # Если точка найдена
                    if found_point:
                        new_point = new_point_from_points_data(
                            found_point, result_parse
                        )
                    else:
                        new_point = new_point_without_city(result_parse)
                        # Если точка не найдена, создаём новую точку

                if new_point is not None:
                    points_data.add_point(new_point) # добавить запись о точке в points_data
                    points_folder.append(new_point)
                    logger.info(f"Добавлена точка: {new_point}")

        # После обработки всех файлов в папке
        if points_folder:
            data_xlsx_path = os.path.join(folder, "data.xlsx")
            save_points_to_excel(points_folder, data_xlsx_path)
            logger.info(f"Точки сохранены в {data_xlsx_path}")

            # Для каждой точки из points_folder создать kml файл с именем
            for point in points_folder:
                # Получаем имя файла без расширения из file_path
                if point.file_path:
                    base_name = os.path.splitext(os.path.basename(point.file_path))[0]
                    kml_file_path = os.path.join(folder, f"{base_name}.kml")
                else:
                    # Если file_path отсутствует, создаем имя на основе координат и времени
                    kml_file_path = os.path.join(folder, f"point_{point.latitude}_{point.longitude}_{point.time.replace(':', '')}.kml")
                create_kml_file(point, kml_file_path)
                logger.info(f"Создан KML файл: {kml_file_path}")

            # Для всех точек в папке создать файл отчета в word
            create_word_report(points_folder, os.path.join(folder, "report.docx"))

    # После обработки всех папок
    # сохранить points_data в файл
    points_data.save()

    # Обработать список wrong_city_data (не найденные города)
    if wrong_city_data:
        # Сохранить в файл для дальнейшего анализа
        wrong_city_file = os.path.join(input_folder, "wrong_cities.txt")
        logger.info("Не найденные города:")
        for city in wrong_city_data:
            logger.info(f" - {city}")
        with open(wrong_city_file, "w", encoding="utf-8") as f:
            f.writelines("\n".join(wrong_city_data))
        logger.info(f"Не найденные города сохранены в {wrong_city_file}")
